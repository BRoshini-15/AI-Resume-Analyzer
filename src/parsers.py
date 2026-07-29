"""
Document Parsers for PDF, DOCX, and Text files.
"""

import io
import re
from typing import Tuple, List
import pypdf
import docx

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts readable text from a PDF file byte stream.
    """
    text = ""
    try:
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        text = f"Error parsing PDF file: {str(e)}"
    return text

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts readable text from a DOCX file byte stream.
    """
    text = ""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table text if present
        for table in doc.tables:
            for row in table.rows:
                paragraphs.extend([cell.text.strip() for cell in row.cells if cell.text.strip()])
        text = "\n".join(paragraphs)
    except Exception as e:
        text = f"Error parsing DOCX file: {str(e)}"
    return text

def parse_uploaded_file(uploaded_file) -> Tuple[str, str]:
    """
    Determines file type and extracts raw text.
    Returns (raw_text, error_message)
    """
    if uploaded_file is None:
        return "", "No file provided"

    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.getvalue()

    if filename.endswith(".pdf"):
        raw_text = extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        raw_text = extract_text_from_docx(file_bytes)
    elif filename.endswith(".txt"):
        raw_text = file_bytes.decode("utf-8", errors="ignore")
    else:
        return "", f"Unsupported file type: {filename}. Please upload PDF, DOCX, or TXT."

    if not raw_text.strip():
        return "", "Extracted text is empty. The file may be an image-only scanned PDF or corrupt."

    return clean_text(raw_text), ""

def clean_text(text: str) -> str:
    """
    Cleans unicode noise, normalizes line breaks and whitespace.
    """
    # Replace non-breaking spaces and special bullet characters
    text = text.replace('\xa0', ' ').replace('\u200b', '')
    text = re.sub(r'[\t\r]+', ' ', text)
    # Normalize multiple line breaks to maximum two
    text = re.sub(r'\n\s*\n', '\n\n', text)
    # Remove multiple spaces
    text = re.sub(r'[ ]{2,}', ' ', text)
    return text.strip()
